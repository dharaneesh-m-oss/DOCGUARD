import os
from docx import Document
from docx.shared import Inches
from typing import List
from .models import UpdateAction

def process_document_update(template_path: str, output_path: str, updates: List[UpdateAction]):
    """
    Loads a .docx template, applies text and image updates, and saves to output_path.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    doc = Document(template_path)

    for action in updates:
        if action.type == "text":
            _replace_text(doc, action.target, action.new_value)
        elif action.type == "image":
            _replace_image(doc, action.target, action.new_value)

    doc.save(output_path)
    return output_path

def _replace_text(doc, target: str, replacement: str):
    """
    Simple text replacement in paragraphs and tables.
    """
    for paragraph in doc.paragraphs:
        if target in paragraph.text:
            paragraph.text = paragraph.text.replace(target, replacement)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if target in paragraph.text:
                        paragraph.text = paragraph.text.replace(target, replacement)

def _replace_image(doc, placeholder_text: str, image_path: str):
    """
    Finds a placeholder text and inserts an image after it.
    Alternatively, could replace existing images if identified by description.
    For this prototype: replaces text 'SIDE_VIEW_PLACEHOLDER' with actual image.
    """
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder_text, "")
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4.0))

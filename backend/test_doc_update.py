import os
from docx import Document
from app.document_service import process_document_update
from app.models import UpdateAction

def create_test_template():
    if not os.path.exists("templates"):
        os.makedirs("templates")
    path = "templates/template.docx"
    doc = Document()
    doc.add_heading('Component Specification Review', 0)
    doc.add_paragraph('Status: {{STATUS}}')
    doc.add_paragraph('Bracket Check: {{BRACKET_CHECK}}')
    doc.add_paragraph('Side view image below:')
    doc.add_paragraph('SIDE_VIEW_PLACEHOLDER')
    doc.save(path)
    print(f"Created test template at {path}")
    return path

def run_test_update():
    template_path = "templates/template.docx"
    output_path = "audit_out/test_updated.docx"
    
    # Mock some updates
    updates = [
        UpdateAction(type="text", target="{{STATUS}}", new_value="Side View Verified (TEST)"),
        UpdateAction(type="text", target="{{BRACKET_CHECK}}", new_value="Verified per spec (TEST)"),
        UpdateAction(type="image", target="SIDE_VIEW_PLACEHOLDER", new_value="test_photo.jpg")
    ]

    # Create a dummy image for testing
    from PIL import Image
    img = Image.new('RGB', (100, 100), color = 'red')
    img.save("test_photo.jpg")

    if not os.path.exists("audit_out"):
        os.makedirs("audit_out")

    try:
        process_document_update(template_path, output_path, updates)
        print(f"Successfully updated document: {output_path}")
    except Exception as e:
        print(f"Error during update: {e}")

if __name__ == "__main__":
    create_test_template()
    run_test_update()
